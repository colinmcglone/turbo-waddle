import pdfquery
import time
import string
import zipfile
import os
import fileinput
import re
import shutil
import sys
from PIL import Image
from PIL import ImageFont
from PIL import ImageDraw
from xml.sax.saxutils import escape
from os.path import basename
from docx import Document

#grabs text from beside known text
def getLeft(key):
	search = pdf.pq('LTTextLineHorizontal:contains("' + key + '")')
	bottom_corner = float(search.attr('y0'))
	right_corner = float(search.attr('x1'))
	top_corner = float(search.attr('y1'))
	answer = pdf.pq('LTTextLineHorizontal:in_bbox("%s, %s, %s, %s")' %(right_corner, bottom_corner, right_corner+210, top_corner+15)).text()
	return answer

#grabs text from under known text
def getUnder(key):
	search = pdf.pq('LTTextLineHorizontal:contains("' + key + '")')
	left_corner = float(search.attr('x0'))
	bottom_corner = float(search.attr('y0'))
	right_corner = float(search.attr('x1'))
	top_corner = float(search.attr('y1'))
	answer = pdf.pq('LTTextLineHorizontal:in_bbox("%s, %s, %s, %s")' %(left_corner, bottom_corner-20, left_corner+150, bottom_corner)).text()
	return answer

#replaces keywords in template with values grabbed from pdf
def replacer_factory(spelling_dict):
    def replacer(match):
        word = match.group()
        return spelling_dict.get(word, word)
    return replacer

def templatefill(text):
    pattern = r'\b\w+\b'  # this pattern matches whole words only
    replacer = replacer_factory(template_values)
    return re.sub(pattern, replacer, text)

def fill():
    with open(documentpath) as in_file:
        text = in_file.read()

    with open(documentpath, 'w') as out_file:
        out_file.write(templatefill(text))

#for handling docx zip
def zipdir(dirPath=None, zipFilePath=None, includeDirInZip=True):
    """Create a zip archive from a directory.
    
    Note that this function is designed to put files in the zip archive with
    either no parent directory or just one parent directory, so it will trim any
    leading directories in the filesystem paths and not include them inside the
    zip archive paths. This is generally the case when you want to just take a
    directory and make it into a zip file that can be extracted in different
    locations. 
    
    Keyword arguments:
    
    dirPath -- string path to the directory to archive. This is the only
    required argument. It can be absolute or relative, but only one or zero
    leading directories will be included in the zip archive.

    zipFilePath -- string path to the output zip file. This can be an absolute
    or relative path. If the zip file already exists, it will be updated. If
    not, it will be created. If you want to replace it from scratch, delete it
    prior to calling this function. (default is computed as dirPath + ".zip")

    includeDirInZip -- boolean indicating whether the top level directory should
    be included in the archive or omitted. (default True)

"""
    if not zipFilePath:
        zipFilePath = dirPath + ".zip"
    if not os.path.isdir(dirPath):
        raise OSError("dirPath argument must point to a directory. "
            "'%s' does not." % dirPath)
    parentDir, dirToZip = os.path.split(dirPath)
	
    #Little nested function to prepare the proper archive path
    def trimPath(path):
        archivePath = path.replace(parentDir, "", 1)
        if parentDir:
            archivePath = archivePath.replace(os.path.sep, "", 1)
        if not includeDirInZip:
            archivePath = archivePath.replace(dirToZip + os.path.sep, "", 1)
        return os.path.normcase(archivePath)
        
    outFile = zipfile.ZipFile(zipFilePath, "w",
        compression=zipfile.ZIP_DEFLATED)
    for (archiveDirPath, dirNames, fileNames) in os.walk(dirPath):
        for fileName in fileNames:
            filePath = os.path.join(archiveDirPath, fileName)
            outFile.write(filePath, trimPath(filePath))
        #Make sure we get empty directories as well
        if not fileNames and not dirNames:
            zipInfo = zipfile.ZipInfo(trimPath(archiveDirPath) + "/")
            outFile.writestr(zipInfo, "")
    outFile.close()

#clear command window
clear = lambda: os.system('cls')
clear()

#get pdf directory from user
pdfd = raw_input("Ensure the required .docx template files are with this script\n\nInput folder with report pdfs: ")

#check path is valid
assert os.path.exists(pdfd), "Files not found at, " + str(pdfd)

#find all pdfs in directory
report = []
for file in os.listdir(pdfd):
	if file.endswith(".pdf"):
		report.append(file)

#used to print number of pdfs converted, number of failed conversions, and which pdfs failed to convert
c = len(report)
s = c
f = []
		
#make directory for finished files
if not os.path.exists(pdfd + "/New Reports and Certificates"):
	os.makedirs(pdfd + "/New Reports and Certificates")
	
#convert each pdf found
for fname in report:
	try:
		#load page 1 of pdf
		pdf = pdfquery.PDFQuery(pdfd + '/' + fname)
		pdf.load(0)
	except:
		print('Error. No file entered, or file not found.\nCheck %s.') %(fname)
		sys.exit(0)
	
	#start grabbing values from pdf for insertion to report
	date = time.strftime("%B %d, %Y").lstrip('0').replace(" 0", " ")
	workorder = pdf.pq('LTTextLineHorizontal:contains("Sales Order")').text()[15:]
	#not all report types have all the same info, so don't throw errors when a grab comes back empty
	try:
		companyname = getUnder("Bill To")
		if str(companyname) == '':
			companyname = getUnder("Ship To")
	except:
		companyname = "N/A"
	
	try:
		search = pdf.pq('LTTextLineHorizontal:contains("Bill To:")')
		left_corner = float(search.attr('x0'))
		bottom_corner = float(search.attr('y0'))
		right_corner = float(search.attr('x1'))
		top_corner = float(search.attr('y1'))
		contactname = pdf.pq('LTTextLineHorizontal:in_bbox("%s, %s, %s, %s")' %(left_corner-10, bottom_corner-30, right_corner+300, bottom_corner-10)).text()
	except:
		contactname = 'NAME'
		
	#Get address block from Ship To
	try:
		search = pdf.pq('LTTextLineHorizontal:contains("Ship To:")')
		left_corner = float(search.attr('x0'))
		bottom_corner = float(search.attr('y0'))
		right_corner = float(search.attr('x1'))
		top_corner = float(search.attr('y1'))
		contact = []
		contact.append(pdf.pq('LTTextLineHorizontal:in_bbox("%s, %s, %s, %s")' %(320, 612, 575, 635)).text())
		contact.append(pdf.pq('LTTextLineHorizontal:in_bbox("%s, %s, %s, %s")' %(320, 582, 575, 605)).text())
		contact.append(pdf.pq('LTTextLineHorizontal:in_bbox("%s, %s, %s, %s")' %(320, 567, 575, 590)).text())
	except:
		contact = ['N/A']
	
	certnumber = pdf.pq('LTTextLineHorizontal:contains("Certificate")').text()[13:]
	
	try:
		make = getLeft("MAKE")
	except:
		make = 'N/A'
	
	try:
		modelnumber = getLeft("MODEL")
	except:
		modelnumber = 'N/A'
	
	try:
		serialnumber = getLeft("SERIAL NO.")
	except:
		serialnumber = 'N/A'
	
	try:
		idnumber = getLeft("ID NO.")
	except:
		idnumber = "N/A"
	
	try:
		location = getLeft("SPECIFIC LOCATION")
	except:
		location = 'N/A'
	
	try:
		lift = getLeft("TYPE")
	except:
		lift = 'N/A'
	
	#Check for faults. the ' (' at the end of most of these is to distinguish the fault lines from other text that may be in the document
	faults = []
	faultlist = ['WIRE ROPE/LOAD LINES (', 
	'CONTROLS/BRAKES/STEERING', 
	'SAFETY DEVICES', 
	'HOOKS/LATCHES/BLOCK (', 
	'HYDRAULIC SYSTEMS (', 
	'OUTRIGGERS/OUTRIGGER BOX (', 
	'TIRES/TRACKS (', 
	'ELECTRICAL APPARATUS (', 
	'BOLTS/NUTS/PINS (', 
	'BOOM/FRAME WELDS (', 
	'STEERING KNUCKLE (', 
	'BOOM FOOT SECTION (', 
	'BOOM HEAD (', 
	'BOOM HOIST (', 
	'SHEAVES (', 
	'ROTATING FRAMES/ BEARINGS (', 
	'HOOK NUT ASSEMBLY (', 
	'DRUM (', 
	'STRUCTURAL ENGINEERING REPORT (', 
	'CAPACITY (', 
	'FRAME (', 
	'LIFTING SURFACE (', 
	'NDE/WELDS (', 
	'HYDRAULIC/VALVES (', 
	'ELECTRICAL/CONTROLS (', 
	'CYLINDER (', 
	'WHEELS (', 
	'BEARINGS (', 
	'LOCKING DEVICE/BRAKES (', 
	'CHAIN (', 
	'FASTENERS/HARDWARE (', 
	'WIRE ROPE/CONNECTORS (', 
	'HOSES/CLAMPS/SEALS (', 
	'MOUNTING BRACKETS (', 
	'PINS/FASTNERS (', 
	'GUARD RAIL/KICK PLATE (', 
	'FALL ARREST ANCHORAGE (', 
	'GATE/CHAIN (', 
	'DIRECTIONAL PLACARD (', 
	'OPERATORS MANUAL (', 
	'LOAD TEST (', 
	'UT TESTING OF PINS (']
	
	#this creates a list of all the faults in the pdf, with the comments
	for i in faultlist:
		fault = ''
		#look for each fault in the pdf one at a time
		e = pdf.pq('LTTextLineHorizontal:contains(' + re.escape(i) + ')')
		#if the fault exists then grab the whole line
		if e.text() is not '':
			fault += e.text()
			left_corner = float(e.attr('x0'))
			bottom_corner = float(e.attr('y0'))
			right_corner = float(e.attr('x1'))
			top_corner = float(e.attr('y1'))
			comment = pdf.pq('LTTextLineHorizontal:in_bbox("%s, %s, %s, %s")' %(right_corner, bottom_corner-15, right_corner+550, top_corner+10)).text()
			fault += ' - ' + comment
			faults.append(fault)

	notok = []
	ok = []
	na = []

	#sort fault lines into categories
	for line in faults:
		if "(NOT OKAY)" in line:
			notok.append(line)
		if "(OKAY)" in line:
			ok.append(line)
		if "(N/A" in line:
			na.append(line)
			
	try:	
		safe = pdf.pq('LTTextLineHorizontal:contains("Status:")').text()[8:]
	except:
		safe = ''
		pass

	#get date of inspection
	try:
		search = pdf.pq('LTTextLineHorizontal:contains("Date:")')
		bottom_corner = float(search.attr('y0'))
		right_corner = float(search.attr('x1'))
		top_corner = float(search.attr('y1'))
		inspectiondate = ''
		inspectiondate = pdf.pq('LTTextLineHorizontal:in_bbox("%s, %s, %s, %s")' %(right_corner-5, bottom_corner-5, right_corner+210, top_corner+15)).text()
		inspectiondate = inspectiondate.lstrip('0').replace(" 0", " ")
	except:
		inspectiondate = ''
		pass
	
	#do it all again for info on second page
	try:
		p2 = pdfquery.PDFQuery(pdfd + '/' + fname)
		p2.load(1)

		search = p2.pq('LTTextLineHorizontal:contains("Inspection Point")')
		left_corner = float(search.attr('x0'))
		bottom_corner = float(search.attr('y0'))
		right_corner = float(search.attr('x1'))
		top_corner = float(search.attr('y1'))

		n = bottom_corner - 10

		faults2 = []
		
		for i in faultlist:
			fault = ''
			e = p2.pq('LTTextLineHorizontal:contains(' + re.escape(i) + ')')
			if e.text() is not '':
				fault += e.text()
				left_corner = float(e.attr('x0'))
				bottom_corner = float(e.attr('y0'))
				right_corner = float(e.attr('x1'))
				top_corner = float(e.attr('y1'))
				comment = p2.pq('LTTextLineHorizontal:in_bbox("%s, %s, %s, %s")' %(right_corner, bottom_corner-15, right_corner+550, top_corner)).text()
				fault += ' - ' + comment
				faults2.append(fault)

		for line in faults2:
			if "(NOT OKAY)" in line:
				notok.append(line)
			if "(OKAY)" in line:
				ok.append(line)
			if "(N/A" in line:
				na.append(line)
				
		search = p2.pq('LTTextLineHorizontal:contains("Date:")')
		bottom_corner = float(search.attr('y0'))
		right_corner = float(search.attr('x1'))
		top_corner = float(search.attr('y1'))
		inspectiondate = ''
		inspectiondate = p2.pq('LTTextLineHorizontal:in_bbox("%s, %s, %s, %s")' %(right_corner-5, bottom_corner-5, right_corner+210, top_corner+15)).text()
		inspectiondate = inspectiondate.lstrip('0').replace(" 0", " ")

		safe = ''
		safe = p2.pq('LTTextLineHorizontal:contains("Status:")').text()[8:]
		
	except:
		print("No second page.")
		pass
		
	#fill in found faults, or no faults found
	if notok == []: #and safe == "SAFE": this was removed because it caused an error when converted to exe <-- Not 100% sure this is accurate
		faultlistf = "1. No faults noted."
	else:
		faultlistf = []
		l = 1
		for each in notok:
			faultlistf.append(str(l) + ". " + escape(each))
			l += 1
		x = '<w:br/><w:br/>'
		faultlistf = x.join(faultlistf)

	#Find lift type, copy template with new name into complete folder
	try:
		machine = getUnder("SAFETY INSPECTION CERTIFICATE")
		if "MOBILE CRANE" in machine:
			newdir = pdfd + '/New Reports and Certificates/' + workorder + ' - ' + companyname + ' - Mobile Crane Report(s)'
			if not os.path.exists(newdir):
				os.makedirs(newdir)
			document = Document('1X-XXXX (XXXXXX) Mobile Crane Report [Template].docx')
			name = workorder + " (" + certnumber + ") Mobile Crane Report.docx"
			document.save(pdfd + '/New Reports and Certificates/' + workorder + ' - ' + companyname + ' - Mobile Crane Report(s)/' + name)
		elif "MAN LIFT" in machine:
			newdir = pdfd + '/New Reports and Certificates/' + workorder + ' - ' + companyname + ' - Man Lift Report(s)'
			if not os.path.exists(newdir):
				os.makedirs(newdir)
			if 'SCISSOR' in lift:
				document = Document('1X-XXXX (XXXXXX) Scissor Lift Report [Template].docx')
				name = workorder + " (" + certnumber + ") Scissor Lift Report.docx"
				document.save(pdfd + '/New Reports and Certificates/' + workorder + ' - ' + companyname + ' - Man Lift Report(s)/' + name)
			if 'BOOM' in lift:
				document = Document('1X-XXXX (XXXXXX) Boom Lift Report [Template].docx')
				name = workorder + " (" + certnumber + ") Boom Lift Report.docx"
				document.save(pdfd + '/New Reports and Certificates/' + workorder + ' - ' + companyname + ' - Man Lift Report(s)/' + name)
			if 'PERSONNEL' in lift:
				document = Document('1X-XXXX (XXXXXX) Personnel Lift Report [Template].docx')
				name = workorder + " (" + certnumber + ") Personnel Lift Report.docx"
				document.save(pdfd + '/New Reports and Certificates/' + workorder + ' - ' + companyname + ' - Man Lift Report(s)/' + name)
		#extract docx components for editing
		with zipfile.ZipFile(newdir + '\\' + name, 'r') as z:
			z.extractall(pdfd + "\\New Reports and Certificates\\tmp\\")
		#path to docx content
		documentpath = pdfd + "/New Reports and Certificates/tmp/word/document.xml"
		#Edit this dict to change replacement keywords, or add new ones. escape() is used to make the text xml safe since the docx content is an xml file
		template_values = {
		'ddate':date, 
		'workorder':escape(workorder), 
		'companyname':escape(companyname), 
		'contactname':contactname.title(),
		'Ccontact1':escape(contact[0]), 
		'Ccontact2':escape(contact[1]), 
		'Ccontact3':escape(contact[2]), 
		'certnumber':escape(certnumber), 
		'mmake':escape(make), 
		'modelnumber':escape(modelnumber), 
		'serialnumber':escape(serialnumber), 
		'idnumber':escape(idnumber), 
		'llocation':escape(location), 
		'inspectiondate':escape(inspectiondate), 
		'faultlist':faultlistf
		}
		#replace from dict into docx template
		fill()
		#edit header
		documentpath = pdfd + "/New Reports and Certificates/tmp/word/header1.xml"
		fill()
		#if there are more than one page, those headers must be edited as well
		try:
			documentpath = pdfd + "/New Reports and Certificates/tmp/word/header2.xml"
			fill()
		except:
			pass
			
		try:
			documentpath = pdfd + "/New Reports and Certificates/tmp/word/header3.xml"
			fill()
		except:
			pass			
		#add the date to the stamp image. .gif format is used because its easier with PIL
		img = Image.open(pdfd + '/New Reports and Certificates/tmp/word/media/image3.gif')
		#the image is converted to RGB because of an error in the .Draw() function in PIL
		img = img.convert('RGB')
		draw = ImageDraw.Draw(img)
		#this font file must be located in the same folder as certs.py
		try:
			font = ImageFont.truetype('verdana.ttf', 32)
		except IOError as e:
			print "I/O error({0}): {1}".format(e.errno, e.strerror)
		except:
			print "Unexpected error:", sys.exc_info()[:2]
			raise
		t = time.strftime("%b %d, %Y")
		draw.text((163,157), t, (0,0,255), font=font) #to change the location of the date on the stamp; change the ordered pair in this line
		img.save(pdfd + '/New Reports and Certificates/tmp/word/media/image3.gif')
		zipdir(pdfd + "\\New Reports and Certificates\\tmp", newdir + '\\' + name, False) #this line uses \\ path notation because of how the zipdir function was written, do not change unless you change zipdir()
		shutil.rmtree(pdfd + '/New Reports and Certificates/tmp')
		shutil.copy(pdfd + '/' + fname, newdir + '/' + name + 'Safety Inspection Certificate.pdf')
		#if something isn't being inserted into the report uncomment this to see if it is being scraped
		'''
		print(machine)
		print(date)
		print(workorder)
		print(companyname)
		for i in contact:
			print i
		print(certnumber)
		print(make)
		print(modelnumber)
		print(serialnumber)
		print(idnumber)
		print(location)
		print(safe)
		for i in ok:
			print i
		for i in notok:
			print i
		for i in na:
			print i
		print(inspectiondate)
		'''
		#prints the old filename and new filename for completed report
		print('DONE %s ->\n%s\n') %(fname, name)
	except IOError as e:
		print "I/O error({0}): {1}".format(e.errno, e.strerror)
		pass
	except:
		s -= 1
		f.append(fname)
		print('ERROR WHILE CREATING REPORT, CHECK REPORT FORMAT OR CONTACT TECHNICAL SUPPORT')
		print "Unexpected error:", sys.exc_info()[:2]
		pass

print('\n%s REPORTS PROCESSED/ %s REPORTS SUCESSFULLY CREATED\n\nFailed reports:') %(c, s)
for each in f:
	print(each)
